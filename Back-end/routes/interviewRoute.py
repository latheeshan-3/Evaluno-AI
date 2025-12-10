"""from fastapi import APIRouter, HTTPException
from schemas.interview import (InterviewQnARequest, InterviewQnAResponse)
from services.llm import chain
import json
from database.connection import cv_results


router = APIRouter(prefix="/interview", tags=["interview"])


@router.post("/", response_model=InterviewQnAResponse)
async def generate_interview_qna(request: InterviewQnARequest):
    try:
        # invoke the chain
        parsed = await chain.ainvoke({
            "cv_text": request.cv_text,
            "job_title": request.job_title,
            "job_requirements": request.job_requirements,
            "job_description": request.job_description
        })
        # parsed is already a dict matching InterviewQnAResponse
        items = json.loads(parsed)

        await cv_results.insert_one({
            "user_id": request.user_id,
            "AIResponse": items
        })

        return {"items": items}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Generation failed: {e}")"""
"""
# routes/interviewRoute.py
from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from schemas.interview import InterviewQnAResponse
from services.llm import chain
from database.connection import cv_results
import json
import fitz  # PyMuPDF
import docx
import io

router = APIRouter(prefix="/interview", tags=["interview"])

# PDF text extraction
def extract_text_from_pdf(file: UploadFile) -> str:
    try:
        file_bytes = file.file.read()
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()  # type: ignore
        return text
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF parse error: {e}")

# DOCX text extraction
def extract_text_from_docx(file: UploadFile) -> str:
    try:
        file_bytes = io.BytesIO(file.file.read())
        document = docx.Document(file_bytes)
        return "\n".join([para.text for para in document.paragraphs])
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"DOCX parse error: {e}")

@router.post("/upload", response_model=InterviewQnAResponse)
async def generate_from_uploaded_cv(
    cv_file: UploadFile = File(...),
    user_id: str = Form(...),
    job_title: str = Form(...),
    job_requirements: str = Form(...),
    job_description: str = Form(...)
):
    try:
        filename = cv_file.filename
        if not filename:
            raise HTTPException(status_code=400, detail="No file uploaded")

        if filename.endswith(".pdf"):
            cv_text = extract_text_from_pdf(cv_file)
        elif filename.endswith(".docx"):
            cv_text = extract_text_from_docx(cv_file)
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type")

        parsed = await chain.ainvoke({
            "cv_text": cv_text,
            "job_title": job_title,
            "job_requirements": job_requirements,
            "job_description": job_description
        })

        items = json.loads(parsed)

        await cv_results.insert_one({
            "user_id": user_id,
            "AIResponse": items
        })

        return {"items": items}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Generation failed: {e}")"""

from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from schemas.interview import InterviewQnAResponse
from services.llm import chain
from database.connection import cv_results
import json
import fitz
import docx
import io
import re
import traceback

router = APIRouter(prefix="/interview", tags=["interview"])


# ---- Utility Functions ----

def extract_text_from_pdf_bytes(file_bytes: bytes) -> str:
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text = ""
        for page in doc:
            text += str(page.get_text())  # Modern PyMuPDF extraction
        return text
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF parse error: {e}")


def extract_text_from_docx_bytes(file_bytes: bytes) -> str:
    try:
        file_stream = io.BytesIO(file_bytes)
        document = docx.Document(file_stream)
        return "\n".join([para.text for para in document.paragraphs])
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"DOCX parse error: {e}")


def strip_json_fence(text: str) -> str:
    """Remove ```json ... ``` or ``` ... ``` fences."""
    cleaned = re.sub(r"```(?:json)?(.*?)```", r"\1", text, flags=re.DOTALL)
    return cleaned.strip()


# ---- API Endpoint ----

@router.post("/upload", response_model=InterviewQnAResponse)
async def generate_from_uploaded_cv(
    cv_file: UploadFile = File(...),
    user_id: str = Form(...),
    job_title: str = Form(...),
    job_requirements: str = Form(...),
    job_description: str = Form(...)
):
    try:
        filename = cv_file.filename
        if not filename:
            raise HTTPException(status_code=400, detail="No file uploaded")

        file_bytes = await cv_file.read()

        # Extract text depending on file type
        if filename.lower().endswith(".pdf"):
            cv_text = extract_text_from_pdf_bytes(file_bytes)
        elif filename.lower().endswith(".docx"):
            cv_text = extract_text_from_docx_bytes(file_bytes)
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type")

        # Call LLM
        raw_response = await chain.ainvoke({
            "cv_text": cv_text,
            "job_title": job_title,
            "job_requirements": job_requirements,
            "job_description": job_description
        })

        cleaned = strip_json_fence(raw_response)

        try:
            items = json.loads(cleaned)
        except json.JSONDecodeError:
            raise HTTPException(status_code=500, detail=f"Invalid JSON from LLM: {raw_response}")

        # Save to DB
        await cv_results.insert_one({
            "user_id": user_id,
            "AIResponse": items
        })

        return {"items": items}

    except Exception as e:
        print("Error in /interview/upload:", traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Generation failed: {e}")

